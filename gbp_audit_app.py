import streamlit as st

# ---------------- USER ACCOUNTS ----------------

USERS = {
    "boss": {"password": "leadfinder123", "role": "admin"},
    "bryan": {"password": "bryan2024", "role": "admin"},
    "user1": {"password": "user1pass", "role": "user"},
    "user2": {"password": "user2pass", "role": "user"},
}

def check_login():
    def login_form():
        with st.form("Login"):
            username = st.text_input("Username")
            password = st.text_input("Password", type="password")
            submit = st.form_submit_button("Login")

            if submit:
                if username in USERS and USERS[username]["password"] == password:
                    st.session_state["logged_in"] = True
                    st.session_state["username"] = username
                    st.session_state["role"] = USERS[username]["role"]
                    st.rerun()
                else:
                    st.error("Invalid username or password")

    if "logged_in" not in st.session_state:
        st.session_state["logged_in"] = False

    if not st.session_state["logged_in"]:
        st.title("GBP Audit Login")
        login_form()
        st.stop()

check_login()

# ---------------- IMPORTS ----------------

import os
import re
import json
import time
import tempfile
import pandas as pd
import requests as req
from io import BytesIO
from datetime import datetime
from apify_client import ApifyClient
import gspread

LOGO_URL = "https://fasthippomedia.com/wp-content/uploads/2024/12/SVG-File.png"

try:
    import anthropic
    HAS_ANTHROPIC = True
except ImportError:
    HAS_ANTHROPIC = False

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    from fpdf import FPDF
    HAS_FPDF = True
except ImportError:
    HAS_FPDF = False

# ---------------- PAGE CONFIG ----------------

st.set_page_config(page_title="GBP Competitor's Audit", page_icon="🏢", layout="wide")

st.title("🏢 GBP Competitor's Audit Tool")

# ---------------- SIDEBAR ----------------

st.sidebar.markdown(f"Logged in as: **{st.session_state.get('username', '')}** ({st.session_state.get('role', '')})")
if st.sidebar.button("Logout"):
    for key in list(st.session_state.keys()):
        del st.session_state[key]
    st.rerun()

st.sidebar.divider()

st.sidebar.markdown("[Check Apify Usage & Billing](https://console.apify.com/billing)")
st.sidebar.markdown("[Check Claude AI Billing](https://console.anthropic.com/settings/billing)")
st.sidebar.markdown("[View Audit History (Google Sheets)](https://docs.google.com/spreadsheets/d/14WqY0En1CuLXRO5fvcy6WKzL_KYWxYV6pUf-FdaHqvU/)")

st.sidebar.divider()

# Downloadable instructions
instructions_text = """GBP COMPETITOR'S AUDIT TOOL - User Guide
========================================

STEP 1: Enter Client Information
- Client Business Name (required)
- Client GBP URL - Google Maps link (required)
- Client Website (optional but recommended)
- Target Keyword / Service (e.g. "plumber in Austin TX")

STEP 2: Find Competitors
- Enter a location (e.g. "Plano, Texas")
- Choose search area: City only, Metro area, or Entire state
- Click "Find Top Competitors" to auto-fill from Google Maps rankings
- OR paste competitor GBP URLs manually
- You can select/deselect competitors from the results list

STEP 3: Configure Scraping Options
- Max reviews per business (default: 100)
- Max photos per business (default: 20)
- Website scraping on/off (scrapes services, about pages, etc.)
- Max pages to crawl per website (default: 5)

STEP 4: Run the Audit
- Click "Run Audit" to start
- Phase 1: Apify scrapes all GBP profiles, reviews, photos, and websites
- Phase 2: Claude AI analyzes the data across 7 sections

AUDIT SECTIONS:
1. Client vs Competitor Overview - comparison table
2. Top 7 Ranking Levers - evidence-based ranking factors
3. GBP Product Instructions - step-by-step product setup
4. Competitor Patterns - observations only, no advice
5. Outlier Analysis - anomalies only, no advice
6. Review Framework - review acquisition & response strategy
7. Photo Upload Plan - weekly photo cadence & types

STEP 5: Download & Save
- Reports auto-save to Google Sheets
- Download as PDF, Word (DOCX), Markdown, or Raw JSON
- View past audits from the sidebar

TIPS:
- Each audit uses Apify credits (scraping) + Anthropic credits (AI analysis)
- A full audit costs roughly $0.05-0.15 in Claude API credits
- Reduce max reviews/photos to save Apify credits on test runs
- Past audits can be reloaded from the sidebar without re-scraping
"""

with st.sidebar.expander("**How to use**", expanded=False):
    st.markdown(
        "1. Enter client info + target keyword\n"
        "2. Click **Find Top Competitors** or paste URLs\n"
        "3. Click **Run Audit**\n"
        "4. Apify scrapes profiles, reviews, photos & websites\n"
        "5. Claude AI analyzes across 7 sections\n"
        "6. Download as PDF, Word, or Markdown\n"
        "7. Auto-saved to Google Sheets"
    )
    st.download_button(
        "📄 Download Full Instructions",
        instructions_text.encode("utf-8"),
        "GBP_Audit_Tool_Instructions.txt",
        "text/plain",
        use_container_width=True,
    )

# ---------------- CONFIG ----------------

AUDIT_SPREADSHEET_ID = "14WqY0En1CuLXRO5fvcy6WKzL_KYWxYV6pUf-FdaHqvU"

APIFY_ACTORS = {
    "gbp_profile": "compass/crawler-google-places",
    "gbp_reviews": "compass/google-maps-reviews-scraper",
    "website_content": "apify/website-content-crawler",
}


# ---------------- GOOGLE SHEETS DATABASE ----------------

def get_gsheet_client():
    """Connect to Google Sheets using service account."""
    try:
        creds_dict = dict(st.secrets["gcp_service_account"])
        return gspread.service_account_from_dict(creds_dict)
    except Exception:
        pass
    creds_path = os.getenv("GOOGLE_SERVICE_ACCOUNT_JSON")
    if creds_path and os.path.exists(creds_path):
        return gspread.service_account(filename=creds_path)
    return None


def save_audit_to_sheets(audit_data, sections):
    """Save audit report to Google Sheets as a persistent record."""
    gc = get_gsheet_client()
    if not gc:
        st.warning("Google Sheets not configured — audit not saved to cloud.")
        return None

    try:
        sh = gc.open_by_key(AUDIT_SPREADSHEET_ID)
    except Exception as e:
        st.warning(f"Could not open Google Sheet: {e}")
        return None

    client_name = audit_data.get("client_name", "Unknown")
    keyword = audit_data.get("target_keyword", "")
    timestamp = audit_data.get("timestamp", datetime.now().strftime("%Y-%m-%d %H:%M:%S"))
    tab_name = f"Audit - {client_name[:15]} - {datetime.now().strftime('%m/%d %H:%M')}"
    tab_name = tab_name[:100]

    try:
        # Create a new tab with the full report
        rows = [
            ["GBP Competitor's Audit Report"],
            ["Client", client_name],
            ["Keyword", keyword],
            ["Generated", timestamp],
            ["Competitors", ", ".join(audit_data.get("comp_labels", []))],
            [""],
        ]

        # Add each section
        for section_name, content in sections.items():
            rows.append([f"=== {section_name} ==="])
            # Split content into rows (Google Sheets has a 50000 char cell limit)
            for line in content.split("\n"):
                rows.append([line])
            rows.append([""])

        ws = sh.add_worksheet(title=tab_name, rows=max(len(rows) + 1, 100), cols=5)
        ws.update(rows, value_input_option="RAW")

    except Exception as e:
        st.warning(f"Could not create audit tab: {e}")
        return None

    # Update history/index on the first sheet
    try:
        history_ws = sh.sheet1
        existing = history_ws.get_all_values()
        if not existing:
            history_ws.update("A1:F1", [["Timestamp", "Client", "Keyword", "Sections", "Tab Name", "Status"]])

        next_row = len(existing) + 1
        history_ws.update(
            f"A{next_row}:F{next_row}",
            [[timestamp, client_name, keyword, len(sections), tab_name, "Complete"]],
        )
    except Exception as e:
        st.caption(f"Could not update history log: {e}")

    return tab_name


def load_audit_history():
    """Load list of past audits from Google Sheets."""
    gc = get_gsheet_client()
    if not gc:
        return []

    try:
        sh = gc.open_by_key(AUDIT_SPREADSHEET_ID)
        history_ws = sh.sheet1
        records = history_ws.get_all_values()
        if len(records) <= 1:
            return []
        headers = records[0]
        audits = []
        for row in records[1:]:
            audit = {}
            for i, h in enumerate(headers):
                audit[h] = row[i] if i < len(row) else ""
            audits.append(audit)
        return list(reversed(audits))  # newest first
    except Exception:
        return []


def load_audit_from_sheet(tab_name):
    """Load a saved audit report from a specific Google Sheets tab."""
    gc = get_gsheet_client()
    if not gc:
        return None

    try:
        sh = gc.open_by_key(AUDIT_SPREADSHEET_ID)
        ws = sh.worksheet(tab_name)
        values = ws.get_all_values()

        # Parse the saved report back into sections
        sections = {}
        current_section = None
        current_content = []
        metadata = {}

        for row in values:
            cell = row[0] if row else ""

            if cell.startswith("=== ") and cell.endswith(" ==="):
                # Save previous section
                if current_section:
                    sections[current_section] = "\n".join(current_content).strip()
                current_section = cell.replace("=== ", "").replace(" ===", "")
                current_content = []
            elif current_section:
                current_content.append(cell)
            elif cell == "Client" and len(row) > 1:
                metadata["client_name"] = row[1]
            elif cell == "Keyword" and len(row) > 1:
                metadata["target_keyword"] = row[1]
            elif cell == "Generated" and len(row) > 1:
                metadata["timestamp"] = row[1]

        # Save last section
        if current_section:
            sections[current_section] = "\n".join(current_content).strip()

        return {"metadata": metadata, "sections": sections}
    except Exception as e:
        st.error(f"Could not load audit: {e}")
        return None


# ---------------- LOGO HELPER ----------------

@st.cache_data(ttl=3600)
def _download_logo():
    """Download the Fast Hippo Media logo and cache it."""
    try:
        r = req.get(LOGO_URL, timeout=10)
        if r.status_code == 200:
            return r.content
    except Exception:
        pass
    return None


# ---------------- BRAND COLORS ----------------
# Fast Hippo Media brand palette
BRAND_NAVY = RGBColor(0x03, 0x04, 0x5E) if HAS_DOCX else None       # #03045E
BRAND_BLUE = RGBColor(0x0C, 0x34, 0xCA) if HAS_DOCX else None       # #0C34CA
BRAND_DARK_TEXT = RGBColor(0x2D, 0x2D, 0x2D) if HAS_DOCX else None  # #2D2D2D
BRAND_GRAY = RGBColor(0x55, 0x55, 0x55) if HAS_DOCX else None       # #555555
BRAND_WHITE = RGBColor(0xFF, 0xFF, 0xFF) if HAS_DOCX else None      # #FFFFFF


# ---------------- DOCX EXPORT ----------------

def _style_heading(heading, color, size=None):
    """Apply brand color to a heading."""
    for run in heading.runs:
        run.font.color.rgb = color
        if size:
            run.font.size = size


def generate_docx(audit_data, sections):
    """Generate a branded Word document from audit results."""
    if not HAS_DOCX:
        return None

    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.font.color.rgb = BRAND_DARK_TEXT

    # ---- COVER PAGE ----
    # Add logo
    logo_data = _download_logo()
    if logo_data:
        logo_stream = BytesIO(logo_data)
        logo_para = doc.add_paragraph()
        logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        logo_para.paragraph_format.space_before = Pt(40)
        run = logo_para.add_run()
        run.add_picture(logo_stream, width=Inches(3))
    else:
        # Fallback text if logo can't be downloaded
        spacer = doc.add_paragraph()
        spacer.paragraph_format.space_before = Pt(80)
        branding = doc.add_paragraph()
        branding.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = branding.add_run("FAST HIPPO MEDIA")
        run.font.size = Pt(14)
        run.font.color.rgb = BRAND_BLUE
        run.bold = True
        run.font.name = "Calibri"

    # Title
    title = doc.add_heading("GBP Competitor's Audit Report", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _style_heading(title, BRAND_NAVY, Pt(28))

    # Divider line
    divider = doc.add_paragraph()
    divider.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = divider.add_run("_" * 50)
    run.font.color.rgb = BRAND_BLUE
    run.font.size = Pt(8)

    # Metadata
    client_name = audit_data.get("client_name", "Client")
    keyword = audit_data.get("target_keyword", "")
    timestamp = audit_data.get("timestamp", "")
    competitors = ", ".join(audit_data.get("comp_labels", []))

    meta_items = [
        ("Client", client_name),
        ("Target Keyword", keyword),
        ("Generated", timestamp),
        ("Competitors", competitors),
    ]
    for label, value in meta_items:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(2)
        run_label = p.add_run(f"{label}: ")
        run_label.font.color.rgb = BRAND_GRAY
        run_label.font.size = Pt(11)
        run_label.font.name = "Calibri"
        run_value = p.add_run(value)
        run_value.font.color.rgb = BRAND_NAVY
        run_value.font.size = Pt(11)
        run_value.bold = True
        run_value.font.name = "Calibri"

    doc.add_page_break()

    # ---- TABLE OF CONTENTS ----
    toc_heading = doc.add_heading("Table of Contents", level=1)
    _style_heading(toc_heading, BRAND_NAVY)

    for i, section_name in enumerate(sections.keys(), 1):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        run_num = p.add_run(f"{i}. ")
        run_num.font.color.rgb = BRAND_BLUE
        run_num.bold = True
        run_num.font.size = Pt(12)
        run_name = p.add_run(section_name.split(". ", 1)[-1] if ". " in section_name else section_name)
        run_name.font.color.rgb = BRAND_DARK_TEXT
        run_name.font.size = Pt(12)

    doc.add_page_break()

    # ---- SECTIONS ----
    for section_name, content in sections.items():
        heading = doc.add_heading(section_name, level=1)
        _style_heading(heading, BRAND_NAVY)

        # Add a colored underline after each section heading
        underline = doc.add_paragraph()
        underline.paragraph_format.space_after = Pt(8)
        run = underline.add_run("_" * 60)
        run.font.color.rgb = BRAND_BLUE
        run.font.size = Pt(6)

        if content.startswith("ERROR"):
            p = doc.add_paragraph(content)
            p.runs[0].font.color.rgb = RGBColor(255, 0, 0)
            continue

        # Parse markdown content into Word formatting
        lines = content.split("\n")
        in_table = False
        table_rows = []
        table_header = []

        for line in lines:
            stripped = line.strip()

            if not stripped:
                if in_table and table_rows:
                    _add_table_to_doc(doc, table_header, table_rows)
                    in_table = False
                    table_rows = []
                    table_header = []
                continue

            # Detect markdown table rows
            if "|" in stripped and stripped.startswith("|"):
                cells = [c.strip() for c in stripped.split("|")[1:-1]]
                if all(re.match(r'^[-:]+$', c) for c in cells if c):
                    continue
                if not in_table:
                    in_table = True
                    table_header = cells
                else:
                    table_rows.append(cells)
                continue

            if in_table and table_rows:
                _add_table_to_doc(doc, table_header, table_rows)
                in_table = False
                table_rows = []
                table_header = []

            # Headings with brand colors
            if stripped.startswith("### "):
                h = doc.add_heading(stripped[4:], level=3)
                _style_heading(h, BRAND_BLUE)
            elif stripped.startswith("## "):
                h = doc.add_heading(stripped[3:], level=2)
                _style_heading(h, BRAND_NAVY)
            elif stripped.startswith("# "):
                h = doc.add_heading(stripped[2:], level=1)
                _style_heading(h, BRAND_NAVY)
            elif stripped.startswith("- ") or stripped.startswith("* "):
                text = stripped[2:]
                p = doc.add_paragraph(style="List Bullet")
                _add_formatted_text(p, text)
            elif re.match(r'^\d+\.\s', stripped):
                text = re.sub(r'^\d+\.\s', '', stripped)
                p = doc.add_paragraph(style="List Number")
                _add_formatted_text(p, text)
            else:
                p = doc.add_paragraph()
                _add_formatted_text(p, stripped)

        if in_table and table_rows:
            _add_table_to_doc(doc, table_header, table_rows)

        doc.add_page_break()

    # ---- FOOTER / BRANDING ----
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_para.add_run("Prepared by Fast Hippo Media | fasthippomedia.com")
    run.font.color.rgb = BRAND_GRAY
    run.font.size = Pt(9)
    run.italic = True

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def _add_table_to_doc(doc, headers, rows):
    """Add a branded table to the Word document."""
    if not headers:
        return
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    num_cols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=num_cols)
    table.style = "Table Grid"

    # Style header row with brand navy background
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i] if i < num_cols else None
        if cell:
            # Set background color
            shading = OxmlElement("w:shd")
            shading.set(qn("w:fill"), "03045E")
            cell._tc.get_or_add_tcPr().append(shading)

            cell.text = header
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
                    run.font.size = Pt(9)
                    run.font.color.rgb = BRAND_WHITE
                    run.font.name = "Calibri"

    # Style data rows with alternating colors
    for r_idx, row_data in enumerate(rows):
        for c_idx, cell_text in enumerate(row_data):
            if c_idx < num_cols:
                cell = table.rows[r_idx + 1].cells[c_idx]
                # Alternating row background
                if r_idx % 2 == 0:
                    shading = OxmlElement("w:shd")
                    shading.set(qn("w:fill"), "F0F4FF")
                    cell._tc.get_or_add_tcPr().append(shading)

                cell.text = cell_text
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(9)
                        run.font.color.rgb = BRAND_DARK_TEXT
                        run.font.name = "Calibri"


def _add_hyperlink(paragraph, url, text=None):
    """Add a clickable hyperlink to a paragraph."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    display_text = text or url
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0C34CA")
    rPr.append(color)

    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)

    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "20")
    rPr.append(sz)

    new_run.append(rPr)
    new_run.text = display_text
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def _add_formatted_text(paragraph, text):
    """Add text with markdown bold/italic formatting and clickable URLs."""
    # First split by URLs
    url_pattern = r'(https?://[^\s\)]+)'
    url_parts = re.split(url_pattern, text)

    for url_part in url_parts:
        if re.match(r'^https?://', url_part):
            # This is a URL - make it clickable
            _add_hyperlink(paragraph, url_part.rstrip(".,;:"))
        else:
            # Process bold/italic within non-URL text
            parts = re.split(r'(\*\*.*?\*\*|\*.*?\*)', url_part)
            for part in parts:
                if part.startswith("**") and part.endswith("**"):
                    run = paragraph.add_run(part[2:-2])
                    run.bold = True
                    run.font.color.rgb = BRAND_NAVY
                elif part.startswith("*") and part.endswith("*"):
                    run = paragraph.add_run(part[1:-1])
                    run.italic = True
                else:
                    paragraph.add_run(part)


# ---------------- PDF EXPORT ----------------

def _sanitize_pdf_text(text):
    """Remove or replace characters that Helvetica can't render."""
    if not text:
        return ""
    # Replace common unicode characters with ASCII equivalents
    replacements = {
        "\u2022": "-",    # bullet •
        "\u2013": "-",    # en dash –
        "\u2014": "--",   # em dash —
        "\u2018": "'",    # left single quote '
        "\u2019": "'",    # right single quote '
        "\u201c": '"',    # left double quote "
        "\u201d": '"',    # right double quote "
        "\u2026": "...",  # ellipsis …
        "\u2605": "*",    # star ★
        "\u2606": "*",    # white star ☆
        "\u2713": "[x]",  # check mark ✓
        "\u2714": "[x]",  # heavy check ✔
        "\u2717": "[ ]",  # cross mark ✗
        "\u2718": "[ ]",  # heavy cross ✘
        "\u00b7": "-",    # middle dot ·
        "\u25cf": "-",    # black circle ●
        "\u25cb": "o",    # white circle ○
        "\u00a0": " ",    # non-breaking space
        "\u200b": "",     # zero-width space
    }
    for char, replacement in replacements.items():
        text = text.replace(char, replacement)
    # Remove any remaining non-latin1 characters
    return text.encode("latin-1", errors="replace").decode("latin-1")


def generate_pdf(audit_data, sections):
    """Generate a branded PDF report with Fast Hippo Media colors."""
    if not HAS_FPDF:
        return None

    # Brand colors (RGB)
    NAVY = (3, 4, 94)       # #03045E
    BLUE = (12, 52, 202)    # #0C34CA
    DARK = (45, 45, 45)     # #2D2D2D
    GRAY = (85, 85, 85)     # #555555
    WHITE = (255, 255, 255)

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=20)

    # ---- COVER PAGE ----
    pdf.add_page()

    # Navy header bar
    pdf.set_fill_color(*NAVY)
    pdf.rect(0, 0, 210, 50, "F")

    # Try to add logo in header
    logo_data = _download_logo()
    if logo_data:
        logo_tmp = tempfile.NamedTemporaryFile(suffix=".png", delete=False)
        logo_tmp.write(logo_data)
        logo_tmp.close()
        try:
            pdf.image(logo_tmp.name, x=55, y=8, w=100)
        except Exception:
            pass
        try:
            os.unlink(logo_tmp.name)
        except Exception:
            pass
    else:
        # Fallback text
        pdf.set_font("Helvetica", "B", 18)
        pdf.set_text_color(*WHITE)
        pdf.set_y(12)
        pdf.cell(0, 10, "FAST HIPPO MEDIA", ln=True, align="C")

    pdf.set_font("Helvetica", "", 10)
    pdf.set_text_color(*WHITE)
    pdf.set_y(38)
    pdf.cell(0, 6, "fasthippomedia.com", ln=True, align="C")

    # Report title
    pdf.set_y(70)
    pdf.set_font("Helvetica", "B", 28)
    pdf.set_text_color(*NAVY)
    pdf.cell(0, 15, "GBP Competitor's", ln=True, align="C")
    pdf.cell(0, 15, "Audit Report", ln=True, align="C")

    # Blue accent line
    pdf.set_draw_color(*BLUE)
    pdf.set_line_width(1)
    pdf.line(60, pdf.get_y() + 5, 150, pdf.get_y() + 5)

    # Metadata
    pdf.set_y(pdf.get_y() + 15)
    pdf.set_font("Helvetica", "", 12)

    client_name = audit_data.get("client_name", "Client")
    keyword = audit_data.get("target_keyword", "")
    timestamp = audit_data.get("timestamp", "")
    competitors = ", ".join(audit_data.get("comp_labels", []))

    meta_items = [
        ("Client", client_name),
        ("Target Keyword", keyword),
        ("Generated", timestamp),
        ("Competitors", competitors),
    ]
    for label, value in meta_items:
        pdf.set_text_color(*GRAY)
        pdf.set_font("Helvetica", "", 11)
        pdf.cell(55, 8, _sanitize_pdf_text(f"{label}:"), align="R")
        pdf.set_text_color(*NAVY)
        pdf.set_font("Helvetica", "B", 11)
        pdf.cell(0, 8, _sanitize_pdf_text(f"  {value}"), ln=True)

    # Footer on cover
    pdf.set_y(260)
    pdf.set_font("Helvetica", "I", 9)
    pdf.set_text_color(*GRAY)
    pdf.cell(0, 6, "Confidential - Prepared exclusively for client use", ln=True, align="C")

    # ---- SECTION PAGES ----
    for section_name, content in sections.items():
        pdf.add_page()

        # Section header bar
        pdf.set_fill_color(*NAVY)
        pdf.rect(0, 0, 210, 20, "F")
        pdf.set_font("Helvetica", "B", 12)
        pdf.set_text_color(*WHITE)
        pdf.set_y(5)
        pdf.cell(0, 10, _sanitize_pdf_text(section_name), ln=True, align="C")

        # Reset position below header
        pdf.set_y(28)
        pdf.set_text_color(*DARK)

        if content.startswith("ERROR"):
            pdf.set_font("Helvetica", "", 11)
            pdf.set_text_color(255, 0, 0)
            pdf.multi_cell(0, 6, _sanitize_pdf_text(content))
            continue

        # Parse content
        lines = content.split("\n")
        in_table_header = True
        for line in lines:
            stripped = line.strip()

            if not stripped:
                pdf.cell(0, 3, "", ln=True)
                continue

            # Table separator - skip
            if re.match(r'^[\|\s\-:]+$', stripped) and "|" in stripped:
                in_table_header = False
                continue

            # Table row
            if stripped.startswith("|") and stripped.endswith("|"):
                cells = [c.strip() for c in stripped.split("|")[1:-1]]
                col_width = (pdf.w - 30) / max(len(cells), 1)

                if in_table_header:
                    # Header row - navy background
                    pdf.set_font("Helvetica", "B", 8)
                    pdf.set_fill_color(*NAVY)
                    pdf.set_text_color(*WHITE)
                    for cell_text in cells:
                        display = cell_text[:55] + "..." if len(cell_text) > 55 else cell_text
                        pdf.cell(col_width, 7, _sanitize_pdf_text(display), border=1, fill=True)
                    pdf.ln()
                    in_table_header = False
                else:
                    # Data row
                    pdf.set_font("Helvetica", "", 8)
                    pdf.set_text_color(*DARK)
                    for cell_text in cells:
                        display = cell_text[:55] + "..." if len(cell_text) > 55 else cell_text
                        pdf.cell(col_width, 6, _sanitize_pdf_text(display), border=1)
                    pdf.ln()
                continue

            in_table_header = True  # Reset for next table

            # Headings
            if stripped.startswith("### "):
                pdf.set_font("Helvetica", "B", 11)
                pdf.set_text_color(*BLUE)
                pdf.cell(0, 8, _sanitize_pdf_text(stripped[4:]), ln=True)
                pdf.set_text_color(*DARK)
            elif stripped.startswith("## "):
                pdf.set_font("Helvetica", "B", 13)
                pdf.set_text_color(*NAVY)
                pdf.cell(0, 9, _sanitize_pdf_text(stripped[3:]), ln=True)
                pdf.set_text_color(*DARK)
            elif stripped.startswith("# "):
                pdf.set_font("Helvetica", "B", 14)
                pdf.set_text_color(*NAVY)
                pdf.cell(0, 10, _sanitize_pdf_text(stripped[2:]), ln=True)
                pdf.set_text_color(*DARK)
            # Bullets
            elif stripped.startswith("- ") or stripped.startswith("* "):
                pdf.set_font("Helvetica", "", 10)
                pdf.set_text_color(*DARK)
                text = stripped[2:].replace("**", "")
                pdf.multi_cell(0, 6, _sanitize_pdf_text(f"  - {text}"))
            # Numbered items
            elif re.match(r'^\d+\.\s', stripped):
                pdf.set_font("Helvetica", "", 10)
                pdf.set_text_color(*DARK)
                text = stripped.replace("**", "")
                pdf.multi_cell(0, 6, _sanitize_pdf_text(text))
            # Regular text
            else:
                pdf.set_font("Helvetica", "", 10)
                pdf.set_text_color(*DARK)
                text = stripped.replace("**", "")
                pdf.multi_cell(0, 6, _sanitize_pdf_text(text))

    # ---- BACK PAGE ----
    pdf.add_page()
    pdf.set_fill_color(*NAVY)
    pdf.rect(0, 0, 210, 297, "F")
    pdf.set_y(120)
    pdf.set_font("Helvetica", "B", 20)
    pdf.set_text_color(*WHITE)
    pdf.cell(0, 12, "FAST HIPPO MEDIA", ln=True, align="C")
    pdf.set_font("Helvetica", "", 12)
    pdf.cell(0, 8, "fasthippomedia.com", ln=True, align="C")
    pdf.cell(0, 15, "", ln=True)
    pdf.set_font("Helvetica", "I", 10)
    pdf.set_text_color(145, 170, 239)  # Light periwinkle
    pdf.cell(0, 8, "Helping businesses dominate local search", ln=True, align="C")

    # Output
    buffer = BytesIO()
    pdf.output(buffer)
    buffer.seek(0)
    return buffer


# ---------------- API KEYS ----------------

def get_apify_token():
    token = os.getenv("APIFY_API_TOKEN")
    if not token:
        try:
            token = st.secrets.get("APIFY_API_TOKEN")
        except Exception:
            pass
    return token


def get_anthropic_key():
    key = os.getenv("ANTHROPIC_API_KEY")
    if not key:
        try:
            key = st.secrets.get("ANTHROPIC_API_KEY")
        except Exception:
            pass
    return key


# ---------------- AUTO-FIND COMPETITORS ----------------

def find_top_competitors(keyword, client_name_to_exclude, location="", num_results=10):
    """Search Google Maps for the keyword in a specific location and return top-ranking businesses."""
    token = get_apify_token()
    if not token:
        return [], "Missing APIFY_API_TOKEN"

    # Combine keyword with location for geo-targeted search
    search_query = f"{keyword} in {location}" if location else keyword

    try:
        client = ApifyClient(token)
        run = client.actor(APIFY_ACTORS["gbp_profile"]).call(
            run_input={
                "searchStringsArray": [search_query],
                "maxCrawledPlacesPerSearch": num_results,
                "language": "en",
                "maxReviews": 0,
                "maxImages": 0,
            },
            timeout_secs=120,
        )
        dataset_id = run["defaultDatasetId"]
        items = list(client.dataset(dataset_id).iterate_items())

        # Filter out the client business
        exclude_lower = client_name_to_exclude.strip().lower() if client_name_to_exclude else ""
        competitors = []
        for item in items:
            name = (item.get("title") or "").strip().lower()
            # Skip if it matches the client name (fuzzy: check if either contains the other)
            if exclude_lower and (exclude_lower in name or name in exclude_lower):
                continue
            competitors.append({
                "name": item.get("title", "Unknown"),
                "url": item.get("url", ""),
                "rating": item.get("totalScore", "N/A"),
                "reviews": item.get("reviewsCount", 0),
                "category": item.get("categoryName", ""),
                "address": item.get("address", ""),
            })

        return competitors, None
    except Exception as e:
        return [], str(e)


# ---- Past Audits (sidebar) ----
st.sidebar.divider()
st.sidebar.markdown("**Past Audits**")
past_audits = load_audit_history()
if past_audits:
    for i, audit in enumerate(past_audits[:10]):
        label = f"{audit.get('Client', '?')} — {audit.get('Keyword', '?')} ({audit.get('Timestamp', '')[:10]})"
        if st.sidebar.button(label, key=f"load_audit_{i}", use_container_width=True):
            tab_name = audit.get("Tab Name", "")
            if tab_name:
                loaded = load_audit_from_sheet(tab_name)
                if loaded:
                    st.session_state["audit_sections"] = loaded["sections"]
                    st.session_state["audit_data"] = {
                        "client_name": loaded["metadata"].get("client_name", ""),
                        "target_keyword": loaded["metadata"].get("target_keyword", ""),
                        "timestamp": loaded["metadata"].get("timestamp", ""),
                        "comp_labels": [],
                    }
                    st.rerun()
else:
    st.sidebar.caption("No past audits found")

# ---------------- INPUT FORM ----------------

st.subheader("Client Information")

col1, col2 = st.columns(2)
with col1:
    client_name = st.text_input("Client Business Name", placeholder="e.g. Smith Plumbing Co")
    client_gbp_url = st.text_input("Client GBP URL (Google Maps link)", placeholder="https://www.google.com/maps/place/...")
with col2:
    client_website = st.text_input("Client Website", placeholder="https://www.smithplumbing.com")
    target_keyword = st.text_input("Target Keyword / Service", placeholder="e.g. plumber in Austin TX")

st.subheader("Competitor GBP URLs")

# Auto-find competitors
st.caption("Auto-fill from current Google Maps rankings, or paste URLs manually")
search_col1, search_col2, search_col3 = st.columns([2, 2, 1])
with search_col1:
    comp_search_location = st.text_input(
        "Competitor search location",
        placeholder="e.g. Plano, Texas or Dallas, TX or Texas",
        key="comp_location",
    )
with search_col2:
    comp_search_radius = st.selectbox(
        "Search area",
        ["City only", "Metro area (nearby cities)", "Entire state"],
        key="comp_radius",
    )
with search_col3:
    st.markdown("<br>", unsafe_allow_html=True)  # vertical spacing
    find_comps = st.button("🔍 Find Top Competitors", use_container_width=True)

if find_comps:
    if not target_keyword:
        st.error("Enter a target keyword first (e.g. 'Digital Marketing Agency')")
    elif not comp_search_location:
        st.error("Enter a location to search competitors in (e.g. 'Plano, Texas')")
    else:
        # Build location string based on search area selection
        location = comp_search_location.strip()
        if comp_search_radius == "Metro area (nearby cities)":
            # Add "near" to broaden slightly
            location = f"near {location}"
        elif comp_search_radius == "Entire state":
            # Extract state from input if city is included
            parts = [p.strip() for p in location.replace(",", " ").split() if p.strip()]
            # Use just the last part (likely the state)
            if len(parts) > 1:
                location = parts[-1]

        with st.spinner(f"Searching Google Maps for: {target_keyword} in {location}..."):
            competitors_found, error = find_top_competitors(target_keyword, client_name, location)
            if error:
                st.error(f"Error finding competitors: {error}")
            elif not competitors_found:
                st.warning("No competitors found. Try a different keyword.")
            else:
                st.session_state["found_competitors"] = competitors_found
                # Auto-fill the top 3 URLs
                for i, comp in enumerate(competitors_found[:3]):
                    st.session_state[f"comp{i + 1}"] = comp["url"]

                st.success(f"Found {len(competitors_found)} businesses ranking for \"{target_keyword}\"")

# Show found competitors as a selectable list
if "found_competitors" in st.session_state and st.session_state["found_competitors"]:
    with st.expander(f"📊 Top ranking businesses for \"{target_keyword}\" (click to select different competitors)", expanded=False):
        comps = st.session_state["found_competitors"]
        st.caption("Top 3 are auto-selected. Check/uncheck to change which competitors to audit.")

        selected_comp_urls = []
        for i, comp in enumerate(comps):
            checked = st.checkbox(
                f"**#{i + 1} {comp['name']}** — ⭐ {comp['rating']} ({comp['reviews']} reviews) | {comp['category']} | {comp['address']}",
                value=(i < 3),
                key=f"comp_select_{i}",
            )
            if checked:
                selected_comp_urls.append(comp["url"])

        # Update the session state with selected competitors
        if selected_comp_urls:
            for i in range(3):
                if i < len(selected_comp_urls):
                    st.session_state[f"comp{i + 1}"] = selected_comp_urls[i]
                else:
                    st.session_state[f"comp{i + 1}"] = ""

comp1_url = st.text_input("Competitor 1 GBP URL", placeholder="https://www.google.com/maps/place/...", key="comp1")
comp2_url = st.text_input("Competitor 2 GBP URL", placeholder="https://www.google.com/maps/place/...", key="comp2")
comp3_url = st.text_input("Competitor 3 GBP URL", placeholder="https://www.google.com/maps/place/...", key="comp3")

st.subheader("Scraping Options")
opt_col1, opt_col2 = st.columns(2)
with opt_col1:
    max_reviews = st.number_input("Max reviews to scrape per business", 10, 500, 100)
    max_photos = st.number_input("Max photos to scrape per business", 5, 100, 20)
with opt_col2:
    scrape_websites = st.checkbox("Scrape business websites for services/content", True)
    max_website_pages = st.number_input("Max pages to crawl per website", 1, 20, 5)

run_audit = st.button("🚀 Run Audit", type="primary", use_container_width=True)


# ---------------- SCRAPING FUNCTIONS ----------------

def extract_place_id_from_url(url):
    """Try to extract a usable search term from a Google Maps URL."""
    if not url:
        return None
    # Clean the URL
    url = url.strip()
    return url


def scrape_gbp_profiles(client, urls, status_container):
    """Scrape GBP profile data for a list of Google Maps URLs."""
    profiles = []
    for i, url in enumerate(urls):
        if not url:
            profiles.append(None)
            continue
        status_container.write(f"Scraping profile {i + 1}/{len(urls)}: {url[:80]}...")
        try:
            run = client.actor(APIFY_ACTORS["gbp_profile"]).call(
                run_input={
                    "startUrls": [{"url": url}],
                    "maxCrawledPlacesPerSearch": 1,
                    "language": "en",
                    "maxImages": max_photos,
                    "maxReviews": 0,  # We use dedicated review scraper
                    "includeOpeningHours": True,
                },
                timeout_secs=120,
            )
            dataset_id = run["defaultDatasetId"]
            items = list(client.dataset(dataset_id).iterate_items())
            if items:
                profiles.append(items[0])
            else:
                profiles.append(None)
                status_container.warning(f"No profile data returned for URL {i + 1}")
        except Exception as e:
            status_container.warning(f"Error scraping profile {i + 1}: {e}")
            profiles.append(None)
    return profiles


def scrape_reviews(client, urls, max_rev, status_container):
    """Scrape reviews for each GBP URL using the dedicated reviews actor."""
    all_reviews = []
    for i, url in enumerate(urls):
        if not url:
            all_reviews.append([])
            continue
        status_container.write(f"Scraping reviews {i + 1}/{len(urls)}: {url[:80]}...")
        try:
            run = client.actor(APIFY_ACTORS["gbp_reviews"]).call(
                run_input={
                    "startUrls": [{"url": url}],
                    "maxReviews": int(max_rev),
                    "reviewsSort": "newest",
                },
                timeout_secs=180,
            )
            dataset_id = run["defaultDatasetId"]
            items = list(client.dataset(dataset_id).iterate_items())
            all_reviews.append(items[:max_rev])
            status_container.write(f"  Got {len(items)} reviews for business {i + 1}")
        except Exception as e:
            status_container.warning(f"Error scraping reviews {i + 1}: {e}")
            all_reviews.append([])
    return all_reviews


def scrape_website_content(client, urls, max_pages, status_container):
    """Scrape website content using the website content crawler."""
    all_content = []
    for i, url in enumerate(urls):
        if not url:
            all_content.append("")
            continue
        status_container.write(f"Scraping website {i + 1}/{len(urls)}: {url[:80]}...")
        try:
            run = client.actor(APIFY_ACTORS["website_content"]).call(
                run_input={
                    "startUrls": [{"url": url}],
                    "maxCrawlPages": int(max_pages),
                    "crawlerType": "cheerio",
                },
                timeout_secs=120,
            )
            dataset_id = run["defaultDatasetId"]
            items = list(client.dataset(dataset_id).iterate_items())
            # Combine all page text
            combined = ""
            for item in items:
                page_text = item.get("text", "") or item.get("markdown", "")
                page_url = item.get("url", "")
                if page_text:
                    combined += f"\n\n--- Page: {page_url} ---\n{page_text[:3000]}"
            all_content.append(combined[:15000])  # Cap total content
        except Exception as e:
            status_container.warning(f"Error scraping website {i + 1}: {e}")
            all_content.append("")
    return all_content


# ---------------- DATA FORMATTING ----------------

def format_profile_summary(profile, label="Business"):
    """Format a GBP profile dict into a readable summary."""
    if not profile:
        return f"{label}: No data available"

    lines = [f"**{label}**"]
    fields = {
        "title": "Name",
        "categoryName": "Primary Category",
        "categories": "All Categories",
        "address": "Address",
        "city": "City",
        "state": "State",
        "phone": "Phone",
        "website": "Website",
        "totalScore": "Rating",
        "reviewsCount": "Review Count",
        "description": "Description",
        "openingHours": "Hours",
    }
    for key, display in fields.items():
        val = profile.get(key)
        if val:
            if isinstance(val, list):
                val = ", ".join(str(v) for v in val)
            elif isinstance(val, dict):
                val = json.dumps(val, ensure_ascii=False)
            lines.append(f"- {display}: {val}")

    # Images count
    images = profile.get("imageUrls", []) or profile.get("images", [])
    if images:
        lines.append(f"- Photo count: {len(images)}")

    return "\n".join(lines)


def format_reviews_summary(reviews, label="Business"):
    """Format review data into a summary for AI analysis."""
    if not reviews:
        return f"{label} Reviews: No reviews scraped"

    lines = [f"**{label} Reviews** ({len(reviews)} scraped)"]

    # Star distribution
    stars = {}
    for r in reviews:
        s = r.get("stars") or r.get("reviewRating") or r.get("rating")
        if s:
            s = int(float(s))
            stars[s] = stars.get(s, 0) + 1
    if stars:
        lines.append(f"- Star distribution: {dict(sorted(stars.items(), reverse=True))}")

    # Review recency
    dates = []
    for r in reviews:
        d = r.get("publishedAtDate") or r.get("publishAt") or r.get("date")
        if d:
            dates.append(str(d)[:10])
    if dates:
        dates.sort(reverse=True)
        lines.append(f"- Most recent review: {dates[0]}")
        lines.append(f"- Oldest scraped: {dates[-1]}")

    # Sample review texts (first 10)
    sample_texts = []
    for r in reviews[:10]:
        text = r.get("text") or r.get("reviewText") or r.get("textTranslated") or ""
        if text:
            sample_texts.append(text[:300])
    if sample_texts:
        lines.append("- Sample review texts:")
        for t in sample_texts:
            lines.append(f'  - "{t}"')

    # Owner replies
    replies = []
    for r in reviews[:20]:
        reply = r.get("ownerResponse") or r.get("responseFromOwnerText") or ""
        if reply:
            replies.append(reply[:200])
    if replies:
        lines.append(f"- Owner replies found: {len(replies)} out of {min(20, len(reviews))} checked")
        lines.append("- Sample owner replies:")
        for rep in replies[:5]:
            lines.append(f'  - "{rep}"')

    # Keywords in reviews
    all_text = " ".join(
        (r.get("text") or r.get("reviewText") or r.get("textTranslated") or "")
        for r in reviews
    ).lower()
    lines.append(f"- Total review text length: {len(all_text)} chars")

    return "\n".join(lines)


def format_photos_summary(profile, label="Business"):
    """Format photo data from a profile."""
    if not profile:
        return f"{label} Photos: No data"

    images = profile.get("imageUrls", []) or profile.get("images", [])
    lines = [f"**{label} Photos**"]
    lines.append(f"- Total photos found: {len(images)}")

    if images:
        lines.append(f"- First 5 photo URLs (for type analysis):")
        for img_url in images[:5]:
            lines.append(f"  - {img_url}")

    return "\n".join(lines)


# ---------------- AI ANALYSIS ----------------

def call_claude(api_key, system_prompt, user_prompt, status_container=None):
    """Call Claude API for analysis. Returns the response text."""
    if not HAS_ANTHROPIC:
        return "ERROR: The `anthropic` Python package is not installed. Run: pip install anthropic"

    if not api_key:
        return "ERROR: No ANTHROPIC_API_KEY found. Set it as an environment variable or in Streamlit secrets."

    try:
        client = anthropic.Anthropic(api_key=api_key)
        message = client.messages.create(
            model="claude-sonnet-4-20250514",
            max_tokens=4096,
            system=system_prompt,
            messages=[{"role": "user", "content": user_prompt}],
        )
        return message.content[0].text
    except Exception as e:
        return f"ERROR calling Claude API: {e}"


def run_section_1(api_key, client_profile, client_website_content, comp_profiles, comp_website_contents, comp_labels):
    """Section 1: Client vs Competitor Overview"""
    system = (
        "You are a local SEO analyst. You will receive scraped Google Business Profile data and website content "
        "for a client and their competitors. Compare them in a structured table format. "
        "Extract: business name, address, services offered, cities/areas served, and key selling points. "
        "Be specific and factual based only on the data provided."
    )

    user = "## Client Profile\n"
    user += format_profile_summary(client_profile, "Client")
    user += f"\n\n## Client Website Content\n{client_website_content[:5000]}\n"

    for i, (prof, content, label) in enumerate(zip(comp_profiles, comp_website_contents, comp_labels)):
        user += f"\n## {label} Profile\n"
        user += format_profile_summary(prof, label)
        user += f"\n\n## {label} Website Content\n{content[:5000]}\n"

    user += (
        "\n\n## Task\n"
        "Create a comparison table with columns: Business Name, Address, Services, Cities Served, Key Selling Points. "
        "Include the client and all competitors. Then write a brief summary of how the client stacks up. "
        "Output as markdown."
    )
    return call_claude(api_key, system, user)


def run_section_2(api_key, client_profile, comp_profiles, comp_reviews, comp_labels, keyword):
    """Section 2: Top 7 Ranking Levers"""
    system = (
        "You are a local SEO analyst specializing in Google Maps rankings. "
        "Based ONLY on the observed top-ranking competitors provided, identify the top 7 ranking levers "
        "Google appears to reward for this keyword. Do NOT give generic SEO advice. "
        "Only cite evidence you can see in the competitor data."
    )

    user = f"## Target Keyword: {keyword}\n\n"
    for i, (prof, revs, label) in enumerate(zip(comp_profiles, comp_reviews, comp_labels)):
        user += f"## {label}\n"
        user += format_profile_summary(prof, label) + "\n"
        user += format_reviews_summary(revs, label) + "\n\n"

    user += (
        "\n## Task\n"
        "List the top 7 ranking levers Google appears to reward for this keyword. "
        "Rank them by impact on map rankings (highest to lowest). "
        "For each lever, cite which competitors demonstrate it and how. "
        "Do not give generic SEO advice or recommendations. "
        "Output strictly as a markdown table with columns: Lever, Evidence from competitors, Why it matters for this keyword."
    )
    return call_claude(api_key, system, user)


def run_section_3(api_key, client_profile, client_website_content, keyword):
    """Section 3: GBP Product Instructions"""
    system = (
        "You are a Google Business Profile specialist. Based on the client's business data, "
        "provide exact step-by-step instructions for adding Products to their GBP. "
        "Be specific to their actual services - not generic instructions."
    )

    user = "## Client Profile\n"
    user += format_profile_summary(client_profile, "Client")
    user += f"\n\n## Client Website Content\n{client_website_content[:5000]}\n"
    user += f"\n## Target Keyword: {keyword}\n"

    user += (
        "\n## Task\n"
        "Based on this client's actual services and business type, provide exact instructions for "
        "adding Products to their Google Business Profile. Include:\n"
        "1. Step-by-step process to add products in GBP\n"
        "2. Specific product names they should add (based on their services)\n"
        "3. Recommended descriptions for each product (keyword-optimized for their market)\n"
        "4. Pricing display recommendations\n"
        "5. Photo recommendations for each product\n"
        "Output as markdown with clear numbered steps."
    )
    return call_claude(api_key, system, user)


def run_section_4(api_key, comp_profiles, comp_reviews, comp_labels, keyword):
    """Section 4: Competitor Patterns (observations only, no advice)"""
    system = (
        "You are a local SEO data analyst. You analyze competitor data and report ONLY patterns and similarities. "
        "Do NOT give suggestions, fixes, advice, conclusions, or next steps. "
        "Output observations only."
    )

    user = ""
    for prof, revs, label in zip(comp_profiles, comp_reviews, comp_labels):
        user += f"## {label}\n"
        user += format_profile_summary(prof, label) + "\n"
        user += format_reviews_summary(revs, label) + "\n"
        user += format_photos_summary(prof, label) + "\n\n"

    user += (
        f"\n## Target Keyword: {keyword}\n\n"
        "## Task\n"
        "Using the data above, identify patterns across all competitors. Report ONLY:\n"
        "- Common primary and secondary categories\n"
        "- Typical review count ranges\n"
        "- Photo upload volume patterns\n"
        "- Keyword usage patterns in business names and descriptions\n\n"
        "Output similarities only. No advice, no conclusions, no next steps. Use markdown tables where appropriate."
    )
    return call_claude(api_key, system, user)


def run_section_5(api_key, comp_profiles, comp_reviews, comp_labels, keyword):
    """Section 5: Outlier Analysis (observations only, no advice)"""
    system = (
        "You are a local SEO data analyst. Identify outliers and anomalies in competitor data. "
        "Do NOT give recommendations, fixes, or advice. Output observations only."
    )

    user = ""
    for prof, revs, label in zip(comp_profiles, comp_reviews, comp_labels):
        user += f"## {label}\n"
        user += format_profile_summary(prof, label) + "\n"
        user += format_reviews_summary(revs, label) + "\n\n"

    user += (
        f"\n## Target Keyword: {keyword}\n\n"
        "## Task\n"
        "Identify clear outliers without giving recommendations:\n"
        "- Which businesses rank high with fewer reviews?\n"
        "- Which rank despite weak branding?\n"
        "- Which ranking factors appear most dominant for this keyword: "
        "proximity, review authority, category relevance, keyword usage, or profile activity/freshness?\n\n"
        "Output observations only. No advice, no fixes. Use markdown."
    )
    return call_claude(api_key, system, user)


def run_section_6(api_key, client_profile, client_reviews, comp_profiles, comp_reviews, comp_labels, client_name, keyword):
    """Section 6: Review Acquisition & Response Framework"""
    system = (
        "You are a local SEO specialist focused on review strategy. "
        "Analyze competitor review data and create an actionable review framework. "
        "Be specific and non-generic. Base everything on the actual competitor data provided."
    )

    user = "## Client\n"
    user += format_profile_summary(client_profile, client_name) + "\n"
    user += format_reviews_summary(client_reviews, client_name) + "\n\n"

    for prof, revs, label in zip(comp_profiles, comp_reviews, comp_labels):
        user += f"## {label}\n"
        user += format_profile_summary(prof, label) + "\n"
        user += format_reviews_summary(revs, label) + "\n\n"

    user += (
        f"\n## Target Keyword: {keyword}\n\n"
        "## Task\n"
        "Analyze the competitor reviews and create a review acquisition and response framework "
        f"for {client_name}. Include:\n"
        "1. Review volume and velocity analysis (how many reviews competitors get and how fast)\n"
        "2. Star distribution comparison\n"
        "3. Recurring keywords in reviews related to services, locations, problems, and outcomes\n"
        "4. How keywords appear organically in reviews and owner replies\n"
        "5. Specific keyword themes to reinforce in review responses\n"
        "6. Review pacing targets (how many per week/month)\n"
        "7. Owner reply patterns and templates (based on what competitors do)\n\n"
        "Make it directly actionable. Avoid vague advice. Use markdown."
    )
    return call_claude(api_key, system, user)


def run_section_7(api_key, client_profile, comp_profiles, comp_labels, client_name, keyword):
    """Section 7: Photo Upload Plan"""
    system = (
        "You are a local SEO specialist focused on GBP photo optimization. "
        "Analyze competitor photo data and create an actionable photo plan. "
        "Be specific and non-generic."
    )

    user = "## Client\n"
    user += format_profile_summary(client_profile, client_name) + "\n"
    user += format_photos_summary(client_profile, client_name) + "\n\n"

    for prof, label in zip(comp_profiles, comp_labels):
        user += f"## {label}\n"
        user += format_profile_summary(prof, label) + "\n"
        user += format_photos_summary(prof, label) + "\n\n"

    user += (
        f"\n## Target Keyword: {keyword}\n\n"
        "## Task\n"
        f"Analyze competitor photos and create a photo upload plan for {client_name} GBP. Include:\n"
        "1. Photo volume comparison across competitors\n"
        "2. Likely photo types based on URLs and business type (job-site, team, exterior, interior, before/after, branded)\n"
        "3. Photo type priorities for the client\n"
        "4. Weekly upload cadence recommendation\n"
        "5. Specific photo ideas based on the client's services\n"
        "6. Geo-tagging and naming recommendations\n\n"
        "Make the plan directly actionable. Avoid vague advice. Use markdown."
    )
    return call_claude(api_key, system, user)


# ---------------- MAIN AUDIT LOGIC ----------------

if run_audit:
    # Validate inputs
    if not client_name:
        st.error("Please enter the client business name.")
        st.stop()
    if not client_gbp_url:
        st.error("Please enter the client GBP URL.")
        st.stop()

    competitor_urls = [u.strip() for u in [comp1_url, comp2_url, comp3_url] if u.strip()]
    if not competitor_urls:
        st.error("Please enter at least one competitor GBP URL.")
        st.stop()

    apify_token = get_apify_token()
    if not apify_token:
        st.error("Missing APIFY_API_TOKEN. Set it as an environment variable or in Streamlit secrets.")
        st.stop()

    anthropic_key = get_anthropic_key()
    if not anthropic_key:
        st.error(
            "Missing ANTHROPIC_API_KEY. Set it as an environment variable or add to .streamlit/secrets.toml:\n\n"
            '`ANTHROPIC_API_KEY = "sk-ant-..."`'
        )
        st.stop()

    apify = ApifyClient(apify_token)

    # Proper capitalization for client name
    client_name = client_name.strip().title()

    # All URLs to scrape
    all_gbp_urls = [client_gbp_url] + competitor_urls

    # Website URLs
    all_website_urls = [client_website]

    # ---- PHASE 1: SCRAPING ----
    st.header("Phase 1: Data Collection")
    scrape_status = st.container()

    # Stop button
    if "audit_stopped" not in st.session_state:
        st.session_state["audit_stopped"] = False
    stop_col1, stop_col2 = st.columns([4, 1])
    with stop_col2:
        if st.button("Stop Audit", type="secondary", use_container_width=True):
            st.session_state["audit_stopped"] = True
            st.warning("Stopping after current step...")

    with st.spinner("Scraping GBP profiles..."):
        scrape_status.subheader("Scraping GBP Profiles")
        all_profiles = scrape_gbp_profiles(apify, all_gbp_urls, scrape_status)
        client_profile = all_profiles[0]
        comp_profiles = all_profiles[1:]
        scrape_status.success(f"Profiles scraped: {sum(1 for p in all_profiles if p)}/{len(all_profiles)}")

    # Use actual business names from scraped profiles as competitor labels
    comp_labels = []
    for i, prof in enumerate(comp_profiles):
        if prof and prof.get("title"):
            comp_labels.append(prof["title"])
        else:
            comp_labels.append(f"Competitor {i + 1}")

    if st.session_state.get("audit_stopped"):
        st.warning("Audit stopped by user after profile scraping.")
        st.session_state["audit_stopped"] = False
        st.stop()

    with st.spinner("Scraping reviews..."):
        scrape_status.subheader("Scraping Reviews")
        all_reviews = scrape_reviews(apify, all_gbp_urls, max_reviews, scrape_status)
        client_reviews = all_reviews[0]
        comp_reviews = all_reviews[1:]
        total_reviews = sum(len(r) for r in all_reviews)
        scrape_status.success(f"Total reviews scraped: {total_reviews}")

    if st.session_state.get("audit_stopped"):
        st.warning("Audit stopped by user after review scraping.")
        st.session_state["audit_stopped"] = False
        st.stop()

    # Scrape websites
    client_website_content = ""
    comp_website_contents = [""] * len(competitor_urls)

    if scrape_websites:
        with st.spinner("Scraping websites..."):
            scrape_status.subheader("Scraping Websites")
            # Gather website URLs from profiles
            website_urls = [client_website or ""]
            for prof in comp_profiles:
                if prof:
                    website_urls.append(prof.get("website", "") or "")
                else:
                    website_urls.append("")

            website_urls_to_scrape = [u for u in website_urls if u]
            if website_urls_to_scrape:
                all_website_content = scrape_website_content(
                    apify, website_urls, max_website_pages, scrape_status
                )
                client_website_content = all_website_content[0] if all_website_content else ""
                comp_website_contents = all_website_content[1:] if len(all_website_content) > 1 else [""] * len(competitor_urls)
                scrape_status.success("Website content scraped")

    if st.session_state.get("audit_stopped"):
        st.warning("Audit stopped by user after website scraping.")
        st.session_state["audit_stopped"] = False
        st.stop()

    # Save raw data to session state
    st.session_state["audit_data"] = {
        "client_name": client_name,
        "client_profile": client_profile,
        "client_reviews": client_reviews,
        "client_website_content": client_website_content,
        "comp_profiles": comp_profiles,
        "comp_reviews": comp_reviews,
        "comp_website_contents": comp_website_contents,
        "comp_labels": comp_labels,
        "target_keyword": target_keyword,
        "timestamp": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
    }

    # ---- PHASE 2: AI ANALYSIS ----
    st.header("Phase 2: AI Analysis")
    analysis_status = st.container()

    sections = {}

    section_runners = [
        ("1. Client vs Competitor Overview", lambda: run_section_1(
            anthropic_key, client_profile, client_website_content,
            comp_profiles, comp_website_contents, comp_labels
        )),
        ("2. Top 7 Ranking Levers", lambda: run_section_2(
            anthropic_key, client_profile, comp_profiles, comp_reviews, comp_labels, target_keyword
        )),
        ("3. GBP Product Instructions", lambda: run_section_3(
            anthropic_key, client_profile, client_website_content, target_keyword
        )),
        ("4. Competitor Patterns", lambda: run_section_4(
            anthropic_key, comp_profiles, comp_reviews, comp_labels, target_keyword
        )),
        ("5. Outlier Analysis", lambda: run_section_5(
            anthropic_key, comp_profiles, comp_reviews, comp_labels, target_keyword
        )),
        ("6. Review Framework", lambda: run_section_6(
            anthropic_key, client_profile, client_reviews, comp_profiles, comp_reviews, comp_labels, client_name, target_keyword
        )),
        ("7. Photo Upload Plan", lambda: run_section_7(
            anthropic_key, client_profile, comp_profiles, comp_labels, client_name, target_keyword
        )),
    ]

    for section_name, runner in section_runners:
        if st.session_state.get("audit_stopped"):
            st.warning(f"Audit stopped by user. Completed {len(sections)} of {len(section_runners)} sections.")
            st.session_state["audit_stopped"] = False
            break
        with st.spinner(f"Analyzing: {section_name}..."):
            analysis_status.write(f"Running: {section_name}")
            result = runner()
            sections[section_name] = result
            if result.startswith("ERROR"):
                analysis_status.error(f"{section_name}: {result}")
            else:
                analysis_status.success(f"{section_name} complete")

    st.session_state["audit_sections"] = sections

    # ---- SAVE TO GOOGLE SHEETS ----
    with st.spinner("Saving audit to Google Sheets..."):
        saved_tab = save_audit_to_sheets(st.session_state["audit_data"], sections)
        if saved_tab:
            st.success(f"Audit saved to Google Sheets tab: **{saved_tab}**")


# ---------------- DISPLAY RESULTS ----------------

if "audit_sections" in st.session_state:
    sections = st.session_state["audit_sections"]
    audit_data = st.session_state.get("audit_data", {})

    st.divider()
    st.header(f"Audit Report: {audit_data.get('client_name', 'Client')}")
    st.caption(f"Target keyword: {audit_data.get('target_keyword', '')} | Generated: {audit_data.get('timestamp', '')}")

    # Create tabs for each section
    tab_names = list(sections.keys())
    tabs = st.tabs(tab_names)

    for tab, (section_name, content) in zip(tabs, sections.items()):
        with tab:
            if content.startswith("ERROR"):
                st.error(content)
            else:
                st.markdown(content)

    # ---- FULL REPORT DOWNLOAD ----
    st.divider()
    st.subheader("Download Report")

    filename_base = f"gbp_audit_{audit_data.get('client_name', 'client').replace(' ', '_').lower()}"

    # Build markdown report
    full_report = f"# GBP Competitor's Audit Report\n"
    full_report += f"**Client:** {audit_data.get('client_name', '')}\n"
    full_report += f"**Keyword:** {audit_data.get('target_keyword', '')}\n"
    full_report += f"**Generated:** {audit_data.get('timestamp', '')}\n"
    full_report += f"**Competitors:** {', '.join(audit_data.get('comp_labels', []))}\n\n"
    full_report += "---\n\n"

    for section_name, content in sections.items():
        full_report += f"## {section_name}\n\n{content}\n\n---\n\n"

    dl_col1, dl_col2, dl_col3, dl_col4 = st.columns(4)

    # PDF
    with dl_col1:
        try:
            pdf_buffer = generate_pdf(audit_data, sections)
        except Exception as pdf_err:
            pdf_buffer = None
            st.caption(f"PDF error: {pdf_err}")
        if pdf_buffer:
            st.download_button(
                "📥 Download PDF",
                pdf_buffer,
                f"{filename_base}.pdf",
                "application/pdf",
                use_container_width=True,
            )
        elif HAS_FPDF:
            st.caption("PDF generation failed — use Word or Markdown instead")
        else:
            st.caption("PDF export unavailable (install fpdf2)")

    # Word DOCX
    with dl_col2:
        docx_buffer = generate_docx(audit_data, sections)
        if docx_buffer:
            st.download_button(
                "📥 Download Word",
                docx_buffer,
                f"{filename_base}.docx",
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
            )
        else:
            st.caption("Word export unavailable (install python-docx)")

    # Markdown
    with dl_col3:
        st.download_button(
            "📥 Download Markdown",
            full_report.encode("utf-8"),
            f"{filename_base}.md",
            "text/markdown",
            use_container_width=True,
        )

    # Raw JSON data
    with dl_col4:
        if audit_data.get("client_profile"):
            raw_data = {
                "client_profile": audit_data.get("client_profile"),
                "competitor_profiles": audit_data.get("comp_profiles"),
                "client_reviews_count": len(audit_data.get("client_reviews", [])),
                "competitor_reviews_counts": [len(r) for r in audit_data.get("comp_reviews", [])],
            }
            st.download_button(
                "📥 Download Raw Data",
                json.dumps(raw_data, indent=2, default=str).encode("utf-8"),
                f"{filename_base}_raw.json",
                "application/json",
                use_container_width=True,
            )
